"""Source-to-Markdown conversion, artifact, Job, and model-role tests."""

import asyncio
from collections.abc import Callable, Iterator
from dataclasses import dataclass
import hashlib
from io import BytesIO
import json
import os
from pathlib import Path
import time
import zipfile

from fastapi.testclient import TestClient
import httpx
from docx import Document
from openpyxl import Workbook
from PIL import Image
import pytest
import respx

from app.config import Settings
from app.db import Base, build_engine
from app.llm.clients import TextGeneration
from app.llm.registry import ModelRoleNotConfiguredError
from app.llm.types import ModelRole, TestedProtocol as Protocol
from app.main import create_app
from app.models.model_config import APIProvider, ModelProfile, ModelRoleBinding
from app.models.source_file import SourceFile
from app.services.document_conversion import (
    ArtifactPublishError,
    DocumentConversionEngine,
    PartDraft,
    SourceDocument,
    VisualConversionModelError,
)
from app.tasks.queue import JobTaskQueue, build_job_task_queue


ADMIN_PASSWORD = "admin-document-conversion-tests"


@dataclass(slots=True)
class ConversionInfra:
    settings: Settings
    application: object
    client: TestClient
    queue: JobTaskQueue


@pytest.fixture
def conversion_infra(tmp_path: Path) -> Iterator[ConversionInfra]:
    settings = Settings(
        admin_password=ADMIN_PASSWORD,
        data_dir=tmp_path / "data",
        source_dir=tmp_path / "sources",
        session_max_age=3600,
    )
    settings.data_dir.mkdir(parents=True)
    settings.source_dir.mkdir(parents=True)
    engine = build_engine(settings)
    Base.metadata.create_all(engine)
    engine.dispose()

    queue = build_job_task_queue(settings, retry_delay=0)
    application = create_app(settings, job_task_queue=queue)
    with TestClient(application) as client:
        login = client.post(
            "/api/auth/admin/login",
            json={"password": ADMIN_PASSWORD},
        )
        assert login.status_code == 200
        yield ConversionInfra(settings, application, client, queue)


def _scan(infra: ConversionInfra) -> list[dict]:
    response = infra.client.post("/api/admin/files/scan")
    assert response.status_code == 200, response.text
    files = infra.client.get("/api/admin/files")
    assert files.status_code == 200
    return files.json()


def _run_next(queue: JobTaskQueue) -> None:
    task = queue.huey.dequeue()
    assert task is not None
    queue.huey.execute(task)


def _png(path: Path) -> None:
    Image.new("RGB", (16, 9), color=(12, 34, 56)).save(path, format="PNG")


def _png_bytes(color: tuple[int, int, int] = (12, 34, 56)) -> bytes:
    stream = BytesIO()
    Image.new("RGB", (16, 9), color=color).save(stream, format="PNG")
    return stream.getvalue()


def test_xlsx_cells_are_deterministic_row_chunks_and_source_is_immutable(
    conversion_infra: ConversionInfra,
) -> None:
    source = conversion_infra.settings.source_dir / "important.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Metrics"
    sheet.append(["name", "amount", "formula"])
    for number in range(1, 205):
        sheet.append([f"row-{number}", number * 1000 + 0.25, f"=B{number + 1}*2"])
    workbook.save(source)
    source_bytes = source.read_bytes()
    source_mtime = source.stat().st_mtime_ns

    record = _scan(conversion_infra)[0]
    queued = conversion_infra.client.post("/api/admin/jobs/convert-changed")
    assert queued.status_code == 202, queued.text
    assert queued.json()["total_items"] == 1

    _run_next(conversion_infra.queue)

    updated = conversion_infra.client.get("/api/admin/files").json()[0]
    artifact_dir = conversion_infra.settings.markdown_dir / str(record["id"])
    manifest = json.loads((artifact_dir / "manifest.json").read_text())
    first = (artifact_dir / "part-001.md").read_text()
    second = (artifact_dir / "part-002.md").read_text()

    assert updated["conversion_status"] == "READY"
    assert len(manifest["parts"]) == 2
    assert manifest["parts"][0]["anchors"] == {
        "rows": "1-200",
        "sheet": "Metrics",
    }
    assert "sheet: \"Metrics\"" in first
    assert "rows: \"1-200\"" in first
    assert "| 2 | row-1 | 1000.25 | =B2*2 |" in first
    assert "| 205 | row-204 | 204000.25 | =B205*2 |" in second
    assert source.read_bytes() == source_bytes
    assert source.stat().st_mtime_ns == source_mtime


def test_xlsx_with_empty_fill_style_is_repaired_in_private_copy(
    conversion_infra: ConversionInfra,
) -> None:
    source = conversion_infra.settings.source_dir / "third-party-export.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Network"
    sheet.append(["service", "address", "port"])
    sheet.append(["vision", "10.0.0.5", 18022])
    workbook.save(source)

    original = BytesIO(source.read_bytes())
    repaired_fixture = BytesIO()
    with zipfile.ZipFile(original) as package, zipfile.ZipFile(
        repaired_fixture,
        "w",
    ) as rewritten:
        for item in package.infolist():
            data = package.read(item.filename)
            if item.filename == "xl/styles.xml":
                valid_fill = b"<fill><patternFill/></fill>"
                assert valid_fill in data
                data = data.replace(valid_fill, b"<fill/>", 1)
            rewritten.writestr(item, data)
    source.write_bytes(repaired_fixture.getvalue())
    source_bytes = source.read_bytes()

    _scan(conversion_infra)
    queued = conversion_infra.client.post("/api/admin/jobs/convert-changed")
    assert queued.status_code == 202
    _run_next(conversion_infra.queue)

    updated = conversion_infra.client.get("/api/admin/files").json()[0]
    markdown = (
        conversion_infra.settings.markdown_dir
        / str(updated["id"])
        / "part-001.md"
    ).read_text()
    assert updated["conversion_status"] == "READY"
    assert "| 2 | vision | 10.0.0.5 | 18022 |" in markdown
    assert source.read_bytes() == source_bytes


def test_convert_changed_incremental_selection_and_explicit_failed_retry(
    conversion_infra: ConversionInfra,
) -> None:
    for name in ("new.txt", "changed.txt", "failed.txt", "ready.txt"):
        (conversion_infra.settings.source_dir / name).write_text(name)
    records = _scan(conversion_infra)
    by_name = {record["filename"]: record for record in records}
    with conversion_infra.application.state.session_factory() as session:
        statuses = {
            "new.txt": "NEW",
            "changed.txt": "CHANGED",
            "failed.txt": "FAILED",
            "ready.txt": "READY",
        }
        for name, status in statuses.items():
            stored = session.get(SourceFile, by_name[name]["id"])
            assert stored is not None
            stored.conversion_status = status
        session.commit()

    first = conversion_infra.client.post("/api/admin/jobs/convert-changed")
    assert first.status_code == 202
    assert {
        item["source_file_id"] for item in first.json()["items"]
    } == {by_name["new.txt"]["id"], by_name["changed.txt"]["id"]}

    retry = conversion_infra.client.post(
        "/api/admin/jobs/convert-changed",
        json={"retry": True},
    )
    assert retry.status_code == 202
    assert [item["source_file_id"] for item in retry.json()["items"]] == [
        by_name["failed.txt"]["id"]
    ]
    assert by_name["ready.txt"]["id"] not in {
        item["source_file_id"]
        for job in (first.json(), retry.json())
        for item in job["items"]
    }


def test_reconvert_all_claims_ready_and_failed_but_not_active_files(
    conversion_infra: ConversionInfra,
) -> None:
    for name in ("ready.txt", "failed.txt", "active.txt"):
        (conversion_infra.settings.source_dir / name).write_text(name)
    records = _scan(conversion_infra)
    by_name = {record["filename"]: record for record in records}
    with conversion_infra.application.state.session_factory() as session:
        for name, status in {
            "ready.txt": "READY",
            "failed.txt": "FAILED",
            "active.txt": "QUEUED",
        }.items():
            stored = session.get(SourceFile, by_name[name]["id"])
            assert stored is not None
            stored.conversion_status = status
        session.commit()

    response = conversion_infra.client.post("/api/admin/jobs/reconvert-all")

    assert response.status_code == 202, response.text
    assert {
        item["source_file_id"] for item in response.json()["items"]
    } == {by_name["ready.txt"]["id"], by_name["failed.txt"]["id"]}
    statuses = {
        record["filename"]: record["conversion_status"]
        for record in conversion_infra.client.get("/api/admin/files").json()
    }
    assert statuses == {
        "active.txt": "QUEUED",
        "failed.txt": "QUEUED",
        "ready.txt": "QUEUED",
    }


def test_batch_conversion_creates_one_job_for_selected_unconverted_files(
    conversion_infra: ConversionInfra,
) -> None:
    for name in (
        "selected-new.txt",
        "selected-failed.txt",
        "not-selected.txt",
        "ready.txt",
    ):
        (conversion_infra.settings.source_dir / name).write_text(name)
    records = _scan(conversion_infra)
    by_name = {record["filename"]: record for record in records}
    with conversion_infra.application.state.session_factory() as session:
        for name, status in {
            "selected-new.txt": "NEW",
            "selected-failed.txt": "FAILED",
            "not-selected.txt": "NEW",
            "ready.txt": "READY",
        }.items():
            stored = session.get(SourceFile, by_name[name]["id"])
            assert stored is not None
            stored.conversion_status = status
        session.commit()

    selected_ids = [
        by_name["selected-new.txt"]["id"],
        by_name["selected-failed.txt"]["id"],
    ]
    response = conversion_infra.client.post(
        "/api/admin/files/batch-convert",
        json={"file_ids": [*selected_ids, selected_ids[0]]},
    )

    assert response.status_code == 202, response.text
    job = response.json()
    assert job["total_items"] == 2
    assert [item["source_file_id"] for item in job["items"]] == selected_ids
    assert conversion_infra.queue.huey.pending_count() == 1
    files = {
        record["filename"]: record
        for record in conversion_infra.client.get("/api/admin/files").json()
    }
    assert files["selected-new.txt"]["conversion_status"] == "QUEUED"
    assert files["selected-failed.txt"]["conversion_status"] == "QUEUED"
    assert files["not-selected.txt"]["conversion_status"] == "NEW"
    assert files["ready.txt"]["conversion_status"] == "READY"

    already_converted = conversion_infra.client.post(
        "/api/admin/files/batch-convert",
        json={"file_ids": [by_name["ready.txt"]["id"]]},
    )
    assert already_converted.status_code == 409


def test_unsupported_is_reported_failed_without_blocking_supported_file(
    conversion_infra: ConversionInfra,
) -> None:
    (conversion_infra.settings.source_dir / "cannot-convert.bin").write_bytes(b"raw")
    (conversion_infra.settings.source_dir / "works.txt").write_text("useful text")
    records = _scan(conversion_infra)
    by_name = {record["filename"]: record for record in records}

    response = conversion_infra.client.post("/api/admin/jobs/convert-changed")
    assert response.status_code == 202
    _run_next(conversion_infra.queue)

    detail = conversion_infra.client.get(
        f"/api/admin/jobs/{response.json()['id']}"
    ).json()
    files = {
        record["filename"]: record
        for record in conversion_infra.client.get("/api/admin/files").json()
    }
    assert detail["status"] == "FAILED"
    assert detail["completed_items"] == 1
    assert detail["failed_items"] == 1
    assert files["cannot-convert.bin"]["conversion_status"] == "UNSUPPORTED"
    assert files["works.txt"]["conversion_status"] == "READY"
    assert not (
        conversion_infra.settings.markdown_dir
        / str(by_name["cannot-convert.bin"]["id"])
    ).exists()
    assert (
        conversion_infra.settings.markdown_dir
        / str(by_name["works.txt"]["id"])
        / "manifest.json"
    ).is_file()


def test_legacy_doc_is_converted_through_libreoffice_bridge(
    conversion_infra: ConversionInfra,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = conversion_infra.settings.source_dir / "旧版方案.doc"
    source.write_bytes(b"legacy-word-placeholder")

    def fake_libreoffice(
        _engine: DocumentConversionEngine,
        _path: Path,
        output_dir: Path,
        heartbeat: Callable[[], None],
    ) -> Path:
        converted = output_dir / "旧版方案.docx"
        document = Document()
        document.add_heading("旧版方案", level=1)
        document.add_paragraph("这是从旧版 Word 转换出的正文。")
        document.save(converted)
        heartbeat()
        return converted

    monkeypatch.setattr(
        DocumentConversionEngine,
        "_convert_legacy_doc_to_docx",
        fake_libreoffice,
    )
    record = _scan(conversion_infra)[0]
    queued = conversion_infra.client.post(
        f"/api/admin/files/{record['id']}/convert"
    )
    assert queued.status_code == 202

    _run_next(conversion_infra.queue)

    updated = conversion_infra.client.get("/api/admin/files").json()[0]
    artifact = conversion_infra.settings.markdown_dir / str(record["id"])
    assert updated["conversion_status"] == "READY"
    assert "旧版方案" in (artifact / "part-001.md").read_text(encoding="utf-8")


def test_conversion_failure_isolated_and_previous_artifact_not_published(
    conversion_infra: ConversionInfra,
) -> None:
    (conversion_infra.settings.source_dir / "bad.json").write_text("{not-json")
    (conversion_infra.settings.source_dir / "good.md").write_text("# Good")
    records = _scan(conversion_infra)
    by_name = {record["filename"]: record for record in records}

    response = conversion_infra.client.post("/api/admin/jobs/convert-changed")
    _run_next(conversion_infra.queue)

    detail = conversion_infra.client.get(
        f"/api/admin/jobs/{response.json()['id']}"
    ).json()
    files = {
        record["filename"]: record
        for record in conversion_infra.client.get("/api/admin/files").json()
    }
    assert detail["status"] == "FAILED"
    assert detail["completed_items"] == 1
    assert detail["failed_items"] == 1
    assert files["bad.json"]["conversion_status"] == "FAILED"
    assert files["good.md"]["conversion_status"] == "READY"
    assert not (
        conversion_infra.settings.markdown_dir / str(by_name["bad.json"]["id"])
    ).exists()
    assert (
        conversion_infra.settings.markdown_dir
        / str(by_name["good.md"]["id"])
        / "part-001.md"
    ).is_file()


def test_atomic_publish_restores_previous_artifact_on_replace_failure(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")
    settings.source_dir.mkdir(parents=True)
    source_path = settings.source_dir / "source.txt"
    source_path.write_text("first artifact")
    source = SourceDocument(
        document_id=7,
        source_path="source.txt",
        source_sha256=hashlib.sha256(source_path.read_bytes()).hexdigest(),
        path=source_path,
    )
    engine = DocumentConversionEngine(
        settings,
        model_resolver=lambda role: pytest.fail(f"unexpected model role: {role}"),
    )
    first = engine.stage(source, job_id=1)
    engine.publish(first)
    destination = settings.markdown_dir / "7"
    previous_manifest = (destination / "manifest.json").read_bytes()
    previous_part = (destination / "part-001.md").read_bytes()

    source_path.write_text("replacement artifact")
    replacement = SourceDocument(
        document_id=7,
        source_path="source.txt",
        source_sha256=hashlib.sha256(source_path.read_bytes()).hexdigest(),
        path=source_path,
    )
    staged = engine.stage(replacement, job_id=2)
    real_replace = os.replace

    def fail_new_directory(source_name, destination_name) -> None:
        if Path(source_name) == staged.staging_dir:
            raise OSError("injected replacement failure")
        real_replace(source_name, destination_name)

    monkeypatch.setattr(
        "app.services.document_conversion.os.replace",
        fail_new_directory,
    )

    with pytest.raises(ArtifactPublishError):
        engine.publish(staged)

    assert (destination / "manifest.json").read_bytes() == previous_manifest
    assert (destination / "part-001.md").read_bytes() == previous_part


def test_staging_directory_shares_markdown_publish_filesystem(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")
    settings.source_dir.mkdir(parents=True)
    source_path = settings.source_dir / "source.txt"
    source_path.write_text("publish safely")
    source = SourceDocument(
        document_id=3,
        source_path=source_path.name,
        source_sha256=hashlib.sha256(source_path.read_bytes()).hexdigest(),
        path=source_path,
    )
    engine = DocumentConversionEngine(
        settings,
        model_resolver=lambda role: pytest.fail(f"unexpected model role: {role}"),
    )

    staged = engine.stage(source, job_id=9)

    assert staged.staging_dir.is_relative_to(settings.markdown_dir)
    engine.publish(staged)
    assert (settings.markdown_dir / "3" / "manifest.json").is_file()


def test_oversized_anchored_draft_is_split_with_stable_segment_locations(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")
    settings.source_dir.mkdir(parents=True)
    source_path = settings.source_dir / "large-page.pdf"
    source_path.write_bytes(b"placeholder")
    source = SourceDocument(
        document_id=8,
        source_path=source_path.name,
        source_sha256=hashlib.sha256(source_path.read_bytes()).hexdigest(),
        path=source_path,
    )
    engine = DocumentConversionEngine(
        settings,
        model_resolver=lambda role: pytest.fail(f"unexpected model role: {role}"),
        text_chars_per_part=1_000,
    )
    monkeypatch.setattr(
        engine,
        "_extract",
        lambda _path, _extension, _heartbeat: [
            PartDraft(body=("可检索证据段落。\n\n" * 400), anchors={"page": 7})
        ],
    )

    staged = engine.stage(source, job_id=1)
    manifest = json.loads((staged.staging_dir / "manifest.json").read_text())

    assert staged.part_count > 1
    assert manifest["converter_version"] == "document-conversion-v3"
    assert manifest["parts"][0]["anchors"]["page"] == 7
    assert manifest["parts"][0]["anchors"]["segment"].startswith("1/")
    assert manifest["parts"][-1]["anchors"]["segment"].split("/")[0] == str(
        staged.part_count
    )
    assert "segment: \"1/" in (staged.staging_dir / "part-001.md").read_text()


def test_oversized_json_keeps_a_complete_code_fence_in_every_part(
    tmp_path: Path,
) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")
    settings.source_dir.mkdir(parents=True)
    source_path = settings.source_dir / "large.json"
    source_path.write_text(
        json.dumps({f"field-{number}": "证据" * 50 for number in range(80)}),
        encoding="utf-8",
    )
    source = SourceDocument(
        document_id=9,
        source_path=source_path.name,
        source_sha256=hashlib.sha256(source_path.read_bytes()).hexdigest(),
        path=source_path,
    )
    engine = DocumentConversionEngine(
        settings,
        model_resolver=lambda role: pytest.fail(f"unexpected model role: {role}"),
        text_chars_per_part=1_000,
    )

    staged = engine.stage(source, job_id=1)

    assert staged.part_count > 1
    for number in range(1, staged.part_count + 1):
        markdown = (staged.staging_dir / f"part-{number:03d}.md").read_text()
        assert markdown.count("```") == 2
        assert "```json\n" in markdown


def test_pdf_uses_text_directly_and_visions_only_low_text_pages(tmp_path: Path) -> None:
    import pymupdf

    settings = Settings(
        data_dir=tmp_path / "data",
        source_dir=tmp_path / "sources",
        document_visual_concurrency=2,
    )
    settings.source_dir.mkdir(parents=True)
    source_path = settings.source_dir / "mixed.pdf"
    with pymupdf.open() as document:
        text_page = document.new_page()
        text_page.insert_text(
            (72, 72),
            "Deterministic searchable evidence " * 4,
        )
        document.new_page()
        document.save(source_path)

    class VisionClient:
        context_window = 128_000
        max_output_tokens = 4_096
        calls = 0

        async def generate_multimodal(self, *_args, **_kwargs) -> TextGeneration:
            self.calls += 1
            return TextGeneration(
                text="blank page visual evidence",
                protocol=Protocol.RESPONSES,
                latency_ms=1,
            )

    client = VisionClient()
    engine = DocumentConversionEngine(
        settings,
        model_resolver=lambda _role: client,
    )
    source = SourceDocument(
        document_id=1,
        source_path=source_path.name,
        source_sha256=hashlib.sha256(source_path.read_bytes()).hexdigest(),
        path=source_path,
    )

    progress_updates: list[dict[str, object]] = []
    staged = engine.stage(
        source,
        job_id=1,
        progress=lambda progress: progress_updates.append(dict(progress)),
    )
    engine.publish(staged)

    assert staged.part_count == 2
    assert client.calls == 1
    first = (settings.markdown_dir / "1" / "part-001.md").read_text()
    second = (settings.markdown_dir / "1" / "part-002.md").read_text()
    assert "Deterministic searchable evidence" in first
    assert "blank page visual evidence" in second
    final_progress = progress_updates[-1]
    assert final_progress["total_pages"] == 2
    assert final_progress["analyzed_pages"] == 2
    assert final_progress["direct_text_pages"] == 1
    assert final_progress["visual_pages"] == 1
    assert final_progress["visual_pages_completed"] == 1
    assert final_progress["model_requests"] == 1
    assert final_progress["phase"] == "publishing"


def test_visual_enrichment_batches_requests_with_configured_concurrency(
    tmp_path: Path,
) -> None:
    settings = Settings(
        data_dir=tmp_path / "data",
        source_dir=tmp_path / "sources",
        document_visual_concurrency=3,
    )

    class ConcurrentVisionClient:
        context_window = 128_000
        max_output_tokens = 4_096
        active = 0
        peak = 0

        async def generate_multimodal(self, *_args, **_kwargs) -> TextGeneration:
            self.active += 1
            self.peak = max(self.peak, self.active)
            await asyncio.sleep(0.02)
            self.active -= 1
            return TextGeneration(
                text="visual evidence",
                protocol=Protocol.RESPONSES,
                latency_ms=20,
            )

    client = ConcurrentVisionClient()
    engine = DocumentConversionEngine(settings, model_resolver=lambda _role: client)

    results = engine._visual_enrichments(
        [
            (_png_bytes((index, index + 1, index + 2)), "image/png")
            for index in range(7)
        ]
    )

    assert results == ["visual evidence"] * 7
    assert client.peak == 3


def test_visual_enrichment_retries_a_transient_model_failure(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")

    class FlakyVisionClient:
        context_window = 128_000
        max_output_tokens = 4_096
        attempts = 0

        async def generate_multimodal(self, *_args, **_kwargs) -> TextGeneration:
            self.attempts += 1
            if self.attempts == 1:
                raise RuntimeError("temporary provider failure")
            return TextGeneration(
                text="recovered visual evidence",
                protocol=Protocol.RESPONSES,
                latency_ms=1,
            )

    client = FlakyVisionClient()
    engine = DocumentConversionEngine(settings, model_resolver=lambda _role: client)
    monkeypatch.setattr(
        "app.services.document_conversion.VISUAL_RETRY_BASE_DELAY_SECONDS",
        0,
    )

    assert engine._visual_enrichment(_png_bytes(), "image/png") == (
        "recovered visual evidence"
    )
    assert client.attempts == 2


def test_visual_enrichment_downscales_large_payloads_before_model_call(
    tmp_path: Path,
) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")
    source = BytesIO()
    Image.new("RGBA", (3_200, 2_400), (12, 34, 56, 180)).save(
        source,
        format="PNG",
    )

    class CapturingVisionClient:
        context_window = 128_000
        max_output_tokens = 4_096
        image_bytes = b""
        media_type = ""

        async def generate_multimodal(
            self,
            _prompt: str,
            image_bytes: bytes,
            *,
            image_media_type: str,
            **_kwargs,
        ) -> TextGeneration:
            self.image_bytes = image_bytes
            self.media_type = image_media_type
            return TextGeneration(
                text="normalized visual evidence",
                protocol=Protocol.RESPONSES,
                latency_ms=1,
            )

    client = CapturingVisionClient()
    engine = DocumentConversionEngine(settings, model_resolver=lambda _role: client)

    assert engine._visual_enrichment(source.getvalue(), "image/png") == (
        "normalized visual evidence"
    )
    assert client.media_type == "image/jpeg"
    with Image.open(BytesIO(client.image_bytes)) as normalized:
        assert max(normalized.size) == 1_600
        assert normalized.mode == "RGB"


def test_visual_enrichment_degrades_legacy_wmf_without_failing_document(
    tmp_path: Path,
) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")

    class VisionClient:
        context_window = 128_000
        max_output_tokens = 4_096
        calls = 0

        async def generate_multimodal(self, *_args, **_kwargs) -> TextGeneration:
            self.calls += 1
            return TextGeneration(
                text="supported image evidence",
                protocol=Protocol.RESPONSES,
                latency_ms=1,
            )

    client = VisionClient()
    engine = DocumentConversionEngine(settings, model_resolver=lambda _role: client)

    results = engine._visual_enrichments(
        [(_png_bytes(), "image/png"), (b"legacy vector bytes", "image/wmf")]
    )

    assert results[0] == "supported image evidence"
    assert "legacy WMF/EMF" in results[1]
    assert client.calls == 1


def test_visual_enrichment_caches_successes_across_retries(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")

    class VisionClient:
        context_window = 128_000
        max_output_tokens = 4_096
        calls = 0

        async def generate_multimodal(self, *_args, **_kwargs) -> TextGeneration:
            self.calls += 1
            return TextGeneration(
                text="durable cached evidence",
                protocol=Protocol.RESPONSES,
                latency_ms=1,
            )

    client = VisionClient()
    engine = DocumentConversionEngine(settings, model_resolver=lambda _role: client)
    image = _png_bytes()

    assert engine._visual_enrichment(image, "image/png") == "durable cached evidence"
    assert engine._visual_enrichment(image, "image/png") == "durable cached evidence"
    assert client.calls == 1
    assert len(list((settings.markdown_dir / ".visual-cache").rglob("*.md"))) == 1


def test_visual_enrichment_preserves_other_successes_when_one_request_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")
    good = _png_bytes((1, 2, 3))
    bad = _png_bytes((4, 5, 6))

    class SelectivelyFailingClient:
        context_window = 128_000
        max_output_tokens = 4_096
        fail_bad = True
        good_calls = 0
        bad_calls = 0

        async def generate_multimodal(
            self,
            _prompt: str,
            image_bytes: bytes,
            **_kwargs,
        ) -> TextGeneration:
            if image_bytes == bad:
                self.bad_calls += 1
                if self.fail_bad:
                    raise RuntimeError("persistent failure")
                text = "recovered bad evidence"
            else:
                self.good_calls += 1
                text = "cached good evidence"
            return TextGeneration(text=text, protocol=Protocol.RESPONSES, latency_ms=1)

    client = SelectivelyFailingClient()
    engine = DocumentConversionEngine(settings, model_resolver=lambda _role: client)
    monkeypatch.setattr(
        "app.services.document_conversion.VISUAL_RETRY_BASE_DELAY_SECONDS",
        0,
    )

    with pytest.raises(VisualConversionModelError, match="vision model failed"):
        engine._visual_enrichments([(good, "image/png"), (bad, "image/png")])
    assert client.good_calls == 1
    assert client.bad_calls == 3

    client.fail_bad = False
    assert engine._visual_enrichments(
        [(good, "image/png"), (bad, "image/png")]
    ) == ["cached good evidence", "recovered bad evidence"]
    assert client.good_calls == 1
    assert client.bad_calls == 4


def test_embedded_visual_enrichment_keeps_successes_and_marks_failed_remainder(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")
    good = _png_bytes((11, 12, 13))
    bad = _png_bytes((21, 22, 23))

    class PartiallyUnavailableClient:
        context_window = 128_000
        max_output_tokens = 4_096

        async def generate_multimodal(
            self,
            _prompt: str,
            image_bytes: bytes,
            **_kwargs,
        ) -> TextGeneration:
            if image_bytes == bad:
                raise RuntimeError("persistent connection failure")
            return TextGeneration(
                text="searchable successful evidence",
                protocol=Protocol.RESPONSES,
                latency_ms=1,
            )

    engine = DocumentConversionEngine(
        settings,
        model_resolver=lambda _role: PartiallyUnavailableClient(),
    )
    monkeypatch.setattr(
        "app.services.document_conversion.VISUAL_RETRY_BASE_DELAY_SECONDS",
        0,
    )

    results = engine._visual_enrichments(
        [(good, "image/png"), (bad, "image/png")],
        allow_partial_failures=True,
    )

    assert results[0] == "searchable successful evidence"
    assert "vision service could not analyze" in results[1]
    assert engine._last_visual_stats["failed_items"] == 1


def test_pptx_conversion_refreshes_heartbeat_during_blocking_extraction(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")
    settings.source_dir.mkdir(parents=True)
    source_path = settings.source_dir / "long-running.pptx"
    source_path.write_bytes(b"test placeholder")
    source = SourceDocument(
        document_id=1,
        source_path=source_path.name,
        source_sha256=hashlib.sha256(source_path.read_bytes()).hexdigest(),
        path=source_path,
    )
    engine = DocumentConversionEngine(
        settings,
        model_resolver=lambda role: pytest.fail(f"unexpected model role: {role}"),
    )
    monkeypatch.setattr(
        "app.services.document_conversion.HEARTBEAT_INTERVAL_SECONDS",
        0.01,
    )

    def blocking_conversion(_path: Path, *, heartbeat: Callable[[], None]) -> str:
        assert callable(heartbeat)
        time.sleep(0.045)
        return "<!-- Slide number: 1 -->\n\n# Restored slide"

    monkeypatch.setattr(
        engine,
        "_convert_pptx_with_markitdown",
        blocking_conversion,
    )
    heartbeat_times: list[float] = []

    staged = engine.stage(
        source,
        job_id=1,
        heartbeat=lambda: heartbeat_times.append(time.monotonic()),
    )

    assert staged.part_count == 1
    assert len(heartbeat_times) >= 5
    assert heartbeat_times[-1] - heartbeat_times[0] >= 0.04


def test_visual_model_wait_refreshes_heartbeat_on_the_request_event_loop(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")

    class SlowVisionClient:
        context_window = 128_000
        max_output_tokens = 4_096
        request_options: dict[str, object] = {}

        async def generate_multimodal(self, *_args, **kwargs) -> TextGeneration:
            self.request_options = kwargs
            await asyncio.sleep(0.045)
            return TextGeneration(
                text="faithful visual evidence",
                protocol=Protocol.RESPONSES,
                latency_ms=45,
            )

    client = SlowVisionClient()
    engine = DocumentConversionEngine(
        settings,
        model_resolver=lambda role: (
            client
            if role is ModelRole.DOCUMENT_CONVERSION
            else pytest.fail(f"unexpected model role: {role}")
        ),
    )
    monkeypatch.setattr(
        "app.services.document_conversion.HEARTBEAT_INTERVAL_SECONDS",
        0.01,
    )
    heartbeat_times: list[float] = []

    result = engine._visual_enrichment(
        _png_bytes(),
        "image/png",
        lambda: heartbeat_times.append(time.monotonic()),
    )

    assert result == "faithful visual evidence"
    assert client.request_options["max_output_tokens"] == 4_096
    assert client.request_options["reasoning_effort"] == "low"
    assert len(heartbeat_times) >= 3
    assert heartbeat_times[-1] - heartbeat_times[0] >= 0.02


def test_visual_model_request_has_a_total_timeout(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")
    cancelled = False

    class HangingVisionClient:
        context_window = 128_000
        max_output_tokens = 4_096

        async def generate_multimodal(self, *_args, **_kwargs) -> TextGeneration:
            nonlocal cancelled
            try:
                await asyncio.Event().wait()
            finally:
                cancelled = True

    engine = DocumentConversionEngine(
        settings,
        model_resolver=lambda _role: HangingVisionClient(),
    )
    monkeypatch.setattr(
        "app.services.document_conversion.HEARTBEAT_INTERVAL_SECONDS",
        0.005,
    )
    monkeypatch.setattr(
        "app.services.document_conversion.MODEL_REQUEST_TIMEOUT_SECONDS",
        0.02,
    )
    monkeypatch.setattr(
        "app.services.document_conversion.VISUAL_RETRY_BASE_DELAY_SECONDS",
        0,
    )
    heartbeat_times: list[float] = []

    with pytest.raises(VisualConversionModelError, match="vision model failed"):
        engine._visual_enrichment(
            _png_bytes(),
            "image/png",
            lambda: heartbeat_times.append(time.monotonic()),
        )

    assert cancelled is True
    assert len(heartbeat_times) >= 2


def test_image_requires_only_document_conversion_role(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path / "data", source_dir=tmp_path / "sources")
    settings.source_dir.mkdir(parents=True)
    image_path = settings.source_dir / "visual.png"
    _png(image_path)
    requested_roles: list[ModelRole] = []

    def unconfigured(role: ModelRole):
        requested_roles.append(role)
        raise ModelRoleNotConfiguredError("not configured")

    engine = DocumentConversionEngine(settings, model_resolver=unconfigured)
    source = SourceDocument(
        document_id=1,
        source_path="visual.png",
        source_sha256=hashlib.sha256(image_path.read_bytes()).hexdigest(),
        path=image_path,
    )

    with pytest.raises(VisualConversionModelError, match="DOCUMENT_CONVERSION"):
        engine.stage(source, job_id=1)
    assert requested_roles == [ModelRole.DOCUMENT_CONVERSION]
    assert not (settings.markdown_dir / "1").exists()


@respx.mock
def test_conversion_profile_a_is_used_and_answer_profile_b_is_never_called(
    conversion_infra: ConversionInfra,
) -> None:
    with conversion_infra.application.state.session_factory() as session:
        cipher = conversion_infra.application.state.api_key_cipher
        conversion_provider = APIProvider(
            name="Conversion provider A",
            provider_type="openai_compatible",
            base_url="https://conversion-a.example.test/v1",
            encrypted_api_key=cipher.encrypt("sk-conversion-a"),
            protocol_preference="responses",
            extra_headers_json={},
            azure_mode="v1",
            enabled=True,
        )
        answer_provider = APIProvider(
            name="Answer provider B",
            provider_type="openai_compatible",
            base_url="https://answer-b.example.test/v1",
            encrypted_api_key=cipher.encrypt("sk-answer-b"),
            protocol_preference="responses",
            extra_headers_json={},
            azure_mode="v1",
            enabled=True,
        )
        conversion_profile = ModelProfile(
            provider=conversion_provider,
            name="profile A",
            remote_model_name="vision-a",
            supports_text=True,
            supports_vision=True,
            supports_structured_output=False,
            tested_protocol="responses",
            enabled=True,
            extra_request_json={},
        )
        answer_profile = ModelProfile(
            provider=answer_provider,
            name="profile B",
            remote_model_name="answer-b",
            supports_text=True,
            supports_vision=True,
            supports_structured_output=False,
            tested_protocol="responses",
            enabled=True,
            extra_request_json={},
        )
        session.add_all([conversion_profile, answer_profile])
        session.flush()
        session.add_all(
            [
                ModelRoleBinding(
                    role=ModelRole.DOCUMENT_CONVERSION.value,
                    model_profile_id=conversion_profile.id,
                ),
                ModelRoleBinding(
                    role=ModelRole.ANSWER_GENERATION.value,
                    model_profile_id=answer_profile.id,
                ),
            ]
        )
        session.commit()

    conversion_route = respx.post(
        "https://conversion-a.example.test/v1/responses"
    ).mock(
        return_value=httpx.Response(
            200,
            json={"output_text": "Transcription: 42\n\nVisual: a small test image."},
        )
    )
    answer_route = respx.post("https://answer-b.example.test/v1/responses").mock(
        return_value=httpx.Response(500)
    )
    embedded_image = conversion_infra.settings.data_dir / "role-check.png"
    _png(embedded_image)
    document_path = conversion_infra.settings.source_dir / "role-check.docx"
    document = Document()
    document.add_picture(str(embedded_image))
    document.save(document_path)
    record = _scan(conversion_infra)[0]

    queued = conversion_infra.client.post(
        f"/api/admin/files/{record['id']}/convert"
    )
    assert queued.status_code == 202, queued.text
    _run_next(conversion_infra.queue)

    updated = conversion_infra.client.get("/api/admin/files").json()[0]
    artifact = (
        conversion_infra.settings.markdown_dir
        / str(record["id"])
        / "part-001.md"
    ).read_text()
    assert updated["conversion_status"] == "READY"
    assert conversion_route.call_count == 1
    assert answer_route.call_count == 0
    assert "Transcription: 42" in artifact
    assert "[Image OCR]" in artifact
